import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from core import ensure_dirs, error_response, success_response


SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".pdf"}


def _validate_file(file_path: str) -> Path:
    path = Path(file_path).expanduser().resolve()

    if not path.exists():
        raise FileNotFoundError(f"File tidak ditemukan: {path}")

    if not path.is_file():
        raise ValueError(f"Path bukan file: {path}")

    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Format tidak didukung: {path.suffix.lower()}. "
            "Gunakan .docx, .xlsx, atau .pdf."
        )

    return path


def _trim_text(text: str, max_chars: Optional[int]) -> Dict[str, Any]:
    if max_chars is None or max_chars <= 0:
        return {
            "text": text,
            "truncated": False,
            "original_char_count": len(text),
            "returned_char_count": len(text),
        }

    if len(text) <= max_chars:
        return {
            "text": text,
            "truncated": False,
            "original_char_count": len(text),
            "returned_char_count": len(text),
        }

    trimmed = text[:max_chars]

    return {
        "text": trimmed,
        "truncated": True,
        "original_char_count": len(text),
        "returned_char_count": len(trimmed),
    }


def read_docx(file_path: str, max_chars: Optional[int] = 50000) -> Dict[str, Any]:
    tool_name = "read_docx"

    try:
        ensure_dirs()
        path = _validate_file(file_path)

        if path.suffix.lower() != ".docx":
            raise ValueError("File harus berformat .docx")

        from docx import Document

        document = Document(str(path))

        paragraphs: List[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        tables: List[Dict[str, Any]] = []
        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])

            tables.append({
                "table_index": table_index,
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
            })

        combined_parts = []
        if paragraphs:
            combined_parts.append("\n".join(paragraphs))

        if tables:
            table_texts = []
            for table in tables:
                table_texts.append(f"[TABLE {table['table_index']}]")
                for row in table["rows"]:
                    table_texts.append(" | ".join(row))
            combined_parts.append("\n".join(table_texts))

        combined_text = "\n\n".join(combined_parts)
        trimmed = _trim_text(combined_text, max_chars)

        core_properties = document.core_properties

        return success_response(
            tool=tool_name,
            message="DOCX berhasil dibaca",
            extra={
                "file_path": str(path),
                "file_type": ".docx",
                "title": core_properties.title,
                "author": core_properties.author,
                "subject": core_properties.subject,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "paragraphs": paragraphs,
                "tables": tables,
                **trimmed,
            },
        )

    except Exception as exc:
        return error_response(tool_name, exc)


def _cell_to_json_safe(value: Any) -> Any:
    if value is None:
        return None

    if isinstance(value, (str, int, float, bool)):
        return value

    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            pass

    return str(value)


def read_xlsx(
    file_path: str,
    sheet_name: Optional[str] = None,
    max_rows: int = 200,
    max_columns: int = 50,
) -> Dict[str, Any]:
    tool_name = "read_xlsx"

    try:
        ensure_dirs()
        path = _validate_file(file_path)

        if path.suffix.lower() != ".xlsx":
            raise ValueError("File harus berformat .xlsx")

        from openpyxl import load_workbook

        workbook = load_workbook(
            filename=str(path),
            read_only=True,
            data_only=True,
        )

        available_sheets = workbook.sheetnames

        if sheet_name:
            if sheet_name not in available_sheets:
                raise ValueError(
                    f"Sheet tidak ditemukan: {sheet_name}. "
                    f"Sheet tersedia: {available_sheets}"
                )
            selected_sheets = [sheet_name]
        else:
            selected_sheets = available_sheets

        sheets: List[Dict[str, Any]] = []

        for current_sheet_name in selected_sheets:
            worksheet = workbook[current_sheet_name]
            rows: List[List[Any]] = []

            for row_index, row in enumerate(
                worksheet.iter_rows(values_only=True),
                start=1,
            ):
                if row_index > max_rows:
                    break

                safe_row = [
                    _cell_to_json_safe(value)
                    for value in row[:max_columns]
                ]

                rows.append(safe_row)

            sheets.append({
                "sheet_name": current_sheet_name,
                "max_row_reported": worksheet.max_row,
                "max_column_reported": worksheet.max_column,
                "returned_row_count": len(rows),
                "returned_column_limit": max_columns,
                "truncated_rows": worksheet.max_row > max_rows,
                "truncated_columns": worksheet.max_column > max_columns,
                "rows": rows,
            })

        workbook.close()

        return success_response(
            tool=tool_name,
            message="XLSX berhasil dibaca",
            extra={
                "file_path": str(path),
                "file_type": ".xlsx",
                "available_sheets": available_sheets,
                "selected_sheets": selected_sheets,
                "sheet_count": len(available_sheets),
                "sheets": sheets,
            },
        )

    except Exception as exc:
        return error_response(tool_name, exc)


def read_pdf(
    file_path: str,
    page_start: Optional[int] = None,
    page_end: Optional[int] = None,
    max_chars: Optional[int] = 50000,
) -> Dict[str, Any]:
    tool_name = "read_pdf"

    try:
        ensure_dirs()
        path = _validate_file(file_path)

        if path.suffix.lower() != ".pdf":
            raise ValueError("File harus berformat .pdf")

        from pypdf import PdfReader

        reader = PdfReader(str(path))
        total_pages = len(reader.pages)

        start_index = 0 if page_start is None else max(page_start - 1, 0)
        end_index = total_pages if page_end is None else min(page_end, total_pages)

        if start_index >= total_pages:
            raise ValueError(
                f"page_start melebihi jumlah halaman. Total halaman: {total_pages}"
            )

        if end_index <= start_index:
            raise ValueError("Rentang halaman tidak valid.")

        pages: List[Dict[str, Any]] = []
        combined_text_parts: List[str] = []

        for index in range(start_index, end_index):
            page = reader.pages[index]
            text = page.extract_text() or ""

            pages.append({
                "page_number": index + 1,
                "char_count": len(text),
                "text": text,
            })

            combined_text_parts.append(
                f"[PAGE {index + 1}]\n{text}"
            )

        combined_text = "\n\n".join(combined_text_parts)
        trimmed = _trim_text(combined_text, max_chars)

        metadata = reader.metadata or {}

        return success_response(
            tool=tool_name,
            message="PDF berhasil dibaca",
            extra={
                "file_path": str(path),
                "file_type": ".pdf",
                "total_pages": total_pages,
                "page_start": start_index + 1,
                "page_end": end_index,
                "metadata": {
                    "title": getattr(metadata, "title", None),
                    "author": getattr(metadata, "author", None),
                    "subject": getattr(metadata, "subject", None),
                    "creator": getattr(metadata, "creator", None),
                },
                "pages": pages,
                **trimmed,
            },
        )

    except Exception as exc:
        return error_response(tool_name, exc)


def read_document(
    file_path: str,
    *,
    sheet_name: Optional[str] = None,
    max_rows: int = 200,
    max_columns: int = 50,
    page_start: Optional[int] = None,
    page_end: Optional[int] = None,
    max_chars: Optional[int] = 50000,
) -> Dict[str, Any]:
    tool_name = "read_document"

    try:
        path = _validate_file(file_path)
        extension = path.suffix.lower()

        if extension == ".docx":
            return read_docx(str(path), max_chars=max_chars)

        if extension == ".xlsx":
            return read_xlsx(
                str(path),
                sheet_name=sheet_name,
                max_rows=max_rows,
                max_columns=max_columns,
            )

        if extension == ".pdf":
            return read_pdf(
                str(path),
                page_start=page_start,
                page_end=page_end,
                max_chars=max_chars,
            )

        raise ValueError(f"Format tidak didukung: {extension}")

    except Exception as exc:
        return error_response(tool_name, exc)
